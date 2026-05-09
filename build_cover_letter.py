"""
build_cover_letter.py
---------------------
Build the Trees - Structure and Function cover letter for paper #2.
"""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT_PATH = ROOT / "output" / "paper" / "cover_letter_paper2.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)


SENDER = [
    "Junsuk Kang, Ph.D.",
    "Department of Landscape Architecture and Rural Systems Engineering",
    "Seoul National University",
    "1 Gwanak-ro, Gwanak-gu, Seoul 08826, Republic of Korea",
    "junkang@snu.ac.kr  |  +82-2-880-4872",
]

DATE_LINE = date.today().strftime("%B %d, %Y")

ADDRESSEE = [
    "The Editor-in-Chief",
    "Trees - Structure and Function",
    "Springer Nature",
]

SUBJECT = ("Manuscript submission: “Probabilistic Eccentric Decay Geometry: "
           "Monte Carlo Confidence Intervals on the Wind-Fatigue Life of Urban "
           "Street Trees”")

GREETING = "Dear Editor,"


BODY_PARAGRAPHS = [
    # Opening: who, what, why this journal
    ("I am pleased to submit the enclosed manuscript, "
     "“Probabilistic Eccentric Decay Geometry: Monte Carlo Confidence "
     "Intervals on the Wind-Fatigue Life of Urban Street Trees,” for "
     "consideration as an Original Article in *Trees - Structure and Function*. "
     "The manuscript is the second in a planned two-part series; the companion "
     "paper (Kang 2026, hereafter Paper I) is currently in press at *Trees - "
     "Structure and Function* and establishes the deterministic absolute-life "
     "framework that the present paper extends to a probabilistic setting."),

    # Background and motivation
    ("Paper I demonstrated that the realistic failure mode of urban street "
     "trees is *decay-accelerated low-cycle fatigue* rather than the high-cycle "
     "fatigue of sound wood, and identified the Mattheck threshold "
     "r_d/R = 0.7 as the point at which lower-bound fatigue life drops below "
     "1.7 years in Seoul. Paper I, however, retains the classical concentric "
     "hollow-cylinder decay model (Mattheck and Breloer 1994), which assumes "
     "circular cavities centred on the stem axis. Field tomography surveys "
     "consistently show that real cavities are eccentric, irregular, and "
     "frequently elongated. Paper I (Section 4.5, fifth limitation) explicitly "
     "identifies this as a non-conservatism that warrants probabilistic "
     "treatment and recommends the present study as the natural follow-on."),

    # The contribution
    ("This manuscript closes that gap. The principal contributions are:"),
]

NUMBERED = [
    ("a vectorisable, analytically grounded framework for arbitrary eccentric "
     "and elliptical decay cavities, verified against the Mattheck (1994) "
     "concentric formula to machine precision and against a Cartesian grid "
     "integration to better than 0.5 %;"),

    ("a 30-scenario Monte Carlo sweep "
     "(5 decay levels x 3 wind distributions x 2 S-N curves, n = 10,000 each) "
     "anchored to the calibrated absolute-life pipeline of Paper I, "
     "yielding the first probabilistic confidence intervals on fatigue life "
     "for decayed urban street trees;"),

    ("a wind-direction *collapse* result — demonstrated analytically and "
     "numerically — showing that under independent uniform priors on "
     "cavity orientation, the wind-rose distribution does not propagate into "
     "the life-ratio correction, identifying cavity-geometry tomography rather "
     "than site-specific wind data as the priority empirical input;"),

    ("an independent validation against ABAQUS Standard 2026 (CPS4R "
     "plane-stress elements, ~ 15,000 elements per case) for six "
     "representative cross-sections, establishing agreement to within 0.08 %;"),

    ("a policy-ready risk-based inspection matrix for three Korean wind "
     "regimes that is directly actionable by municipal forestry departments."),
]

CONTINUE = [
    # Headline numerical result
    ("The headline numerical result is that, at the Mattheck threshold "
     "r_d/R = 0.7, the *median* eccentric tree has a fatigue life of "
     "25 % of the concentric Paper I estimate (lower-bound S-N), and the "
     "5th-percentile tail is shorter by a factor of 47 to 143 depending on "
     "the S-N curve. This invalidates deterministic point estimates as the "
     "basis for risk-based management and motivates a percentile-driven "
     "inspection policy. We translate the result into a 3-region x 5-decay "
     "policy matrix that is, to our knowledge, the first probabilistically "
     "calibrated tree-risk schedule in the Korean context."),

    # Compatibility with Paper I
    ("The present manuscript is fully compatible with Paper I in three "
     "concrete senses: (i) the closed-form stress-amplification formula "
     "reduces to the Mattheck formula in the concentric limit to machine "
     "precision; (ii) absolute-life anchoring uses Paper I's Table 2 "
     "verbatim, treating Paper I's full pipeline as a calibrated baseline; "
     "and (iii) the eccentric correction acts as a multiplicative stress "
     "factor on the stress history, the same operator Paper I uses for "
     "concentric stress amplification. The two papers therefore form a "
     "logically nested pair: Paper I establishes the absolute-life pipeline, "
     "and the present paper supplies the geometry-uncertainty correction "
     "that maps the deterministic point estimates onto probabilistic "
     "confidence intervals. We respectfully suggest that the two manuscripts "
     "be considered as a planned series in *Trees - Structure and Function*."),

    # Suitability for Trees
    ("The manuscript fits squarely within the *Trees - Structure and "
     "Function* scope: it is centred on tree biomechanics and explicitly "
     "couples experimentally observable cavity-geometry priors with "
     "structural fatigue prediction and arboricultural management practice. "
     "The risk matrix in Section 3.7 and the discussion of the implications "
     "for the Korean Tree Risk Management standard "
     "(Ministry of Land, Infrastructure and Transport 2019) make the "
     "manuscript directly relevant to the journal's professional readership "
     "in addition to its biomechanical-research audience."),

    # Originality and ethics declarations
    ("The manuscript represents original work that is not under "
     "consideration elsewhere. All Monte Carlo sample data, ABAQUS "
     "validation input files, and analysis scripts will be archived on "
     "GitHub at submission. The author has no competing interests to "
     "declare."),

    # Suggested reviewers
    ("If helpful, suggested reviewers with relevant expertise include "
     "(i) Prof. Klaus Mattheck (Karlsruhe Institute of Technology) on "
     "decay-mechanics modelling, (ii) Prof. Kenneth James (University of "
     "Melbourne) on tree dynamics and crown-stem coupling, and "
     "(iii) Dr. Frank Schwarze (EMPA) on wood-decay fungal mechanisms; "
     "we have no co-authored or institutional conflict with any of the above. "
     "The author is happy to suggest additional reviewers at the editor's "
     "request."),

    # Closing
    ("Thank you for considering this manuscript. I look forward to your "
     "response and to working with the editorial team."),
]

SIGN_OFF = [
    "",
    "Sincerely,",
    "",
    "",
    "Junsuk Kang, Ph.D.",
    "Corresponding author",
]


# ─────────────────────────────────────────────────────────────────────
# Build the document
# ─────────────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)


def add_p(text="", italic=False, bold=False, size=12,
          align=WD_ALIGN_PARAGRAPH.LEFT, after=10, space_before=0,
          line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line_spacing
    if text:
        # support inline asterisk-italics for *foo* segments
        chunks = []
        i = 0
        while i < len(text):
            j = text.find("*", i)
            if j < 0:
                chunks.append((text[i:], False))
                break
            if j > i:
                chunks.append((text[i:j], False))
            k = text.find("*", j + 1)
            if k < 0:
                chunks.append((text[j:], False))
                break
            chunks.append((text[j+1:k], True))
            i = k + 1
        for chunk, is_it in chunks:
            r = p.add_run(chunk)
            r.font.name = BODY_FONT
            r.font.size = Pt(size)
            r.font.bold = bold
            r.font.italic = italic or is_it
            r.font.color.rgb = RGBColor(0, 0, 0)
    return p


# 1. Sender block (right-aligned)
for line in SENDER:
    add_p(line, align=WD_ALIGN_PARAGRAPH.RIGHT, after=2, line_spacing=1.0)

add_p("", after=6)
add_p(DATE_LINE, align=WD_ALIGN_PARAGRAPH.RIGHT, after=18, line_spacing=1.0)

# 2. Addressee
for line in ADDRESSEE:
    add_p(line, after=2, line_spacing=1.0)

add_p("", after=6)
# Subject line
add_p(SUBJECT, bold=True, after=14, line_spacing=1.15)

# 3. Greeting
add_p(GREETING, after=12, line_spacing=1.15)

# 4. Body paragraphs
for para in BODY_PARAGRAPHS:
    add_p(para, after=10, line_spacing=1.15)

# Numbered list
for i, item in enumerate(NUMBERED, start=1):
    add_p(f"({i}) {item}", after=8, line_spacing=1.15)

add_p("", after=4)
for para in CONTINUE:
    add_p(para, after=10, line_spacing=1.15)

# 5. Sign-off
for line in SIGN_OFF:
    add_p(line, after=4, line_spacing=1.0)

doc.save(str(OUT_PATH))
print(f"Saved: {OUT_PATH}")
