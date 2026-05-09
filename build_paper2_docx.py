"""
build_paper2_docx.py
--------------------
Build the Trees-Springer DOCX manuscript from paper2_draft_v2.md.

Adapted from fatigue-tree/build_docx.py (paper #1). Differences:
  - Tables are kept inline (pandoc renders the markdown tables in §3 + Appendix A)
  - Figures (11 main + Fig A1) are appended after References as a List of Figures
  - All other formatting (font, double-spacing, margins, black text) is the same

Usage:
  python build_paper2_docx.py
"""

from __future__ import annotations

import os
import re
from pathlib import Path

import pypandoc
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT     = Path(__file__).resolve().parent
MD_PATH  = ROOT / "output" / "paper" / "paper2_draft_v2.md"
OUT_PATH = ROOT / "output" / "paper" / "paper2_Trees.docx"
FIG_DIR  = ROOT / "output" / "figures"
TMP_MD   = ROOT / "_tmp_body_p2.md"
TMP_DOCX = ROOT / "_tmp_body_p2.docx"

BODY_FONT = "Times New Roman"


# ─────────────────────────────────────────────────────────────────────
# Figure list (with provisional captions)
# ─────────────────────────────────────────────────────────────────────
FIGURES = [
    {"file": "fig_regression_test.png",
     "caption": ("Figure 1. Four-stage regression test of the eccentric-decay framework "
                 "against the Mattheck (1994) concentric limit. "
                 "(a) The closed-form SAF reproduces the analytical formula to machine precision "
                 "across decay ratios. "
                 "(b) Polar plot showing direction-invariance of SAF for concentric cavities. "
                 "(c) Continuity of SAF_worst as eccentricity tends to zero. "
                 "(d) Closed-form section properties match a 1500x1500 Cartesian grid integration "
                 "to better than 0.5%.")},
    {"file": "eccentric_demo.png",
     "caption": ("Figure 2. Single-sample SAF behaviour at r_d/R = 0.5 for three representative "
                 "cross-sections. "
                 "(a) Cross-section schematics for concentric, eccentric circular (e_norm = 0.5), "
                 "and area-equivalent elliptical 2:1 cavities, showing the net-section centroid (X). "
                 "(b) Polar SAF plot. "
                 "(c) Worst-case SAF as a function of normalised eccentricity at four decay ratios. "
                 "(d) SAF as a function of decay ratio for the three cross-section types. "
                 "(e) Net-section centroid shift opposite to the cavity offset. "
                 "(f) Worst-direction angle as eccentricity increases.")},
    {"file": "fig_mc_cdf.png",
     "caption": ("Figure 3. Cumulative distribution functions of the life ratio rho at five decay levels "
                 "under three wind distributions (lower-bound S-N, n = 10,000 samples). "
                 "The three wind distributions overlay almost exactly, demonstrating the "
                 "wind-direction collapse result.")},
    {"file": "fig_mc_percentile.png",
     "caption": ("Figure 4. Life ratio rho versus decay ratio r_d/R: percentile bands "
                 "(5-95% and 25-75%) and median for lower-bound and mean S-N curves. "
                 "Vertical dashed line marks the Mattheck threshold r_d/R = 0.7.")},
    {"file": "fig_mc_heatmap.png",
     "caption": ("Figure 5. Life-shortening factor 1/rho heatmap by sample percentile and decay ratio. "
                 "The bottom 5% tail at the Mattheck threshold reaches a 47-fold (lower-bound) or "
                 "143-fold (mean) reduction relative to the concentric assumption.")},
    {"file": "fig_mc_concentric_vs_ecc.png",
     "caption": ("Figure 6. Distribution of the life ratio at the Mattheck threshold r_d/R = 0.7. "
                 "The eccentric Monte Carlo distribution (blue histogram) shows that the deterministic "
                 "concentric estimate (black line, rho = 1) lies far above the median (red dashed) "
                 "and the 5th-percentile (orange dotted) of the population.")},
    {"file": "fig_absolute_distribution_rd07.png",
     "caption": ("Figure 7. Absolute fatigue life distribution at r_d/R = 0.7 (Seoul wind regime, "
                 "ginkgo, H = 8 m, DBH = 20 cm) anchored to Paper I Table 2. "
                 "The concentric estimate (1.7 yr lower-bound, 6.9 yr mean) overpredicts the "
                 "median (5 months) and the 5th-percentile tail (13 days).")},
    {"file": "fig_absolute_percentile.png",
     "caption": ("Figure 8. Percentile bands of absolute fatigue life as a function of decay ratio. "
                 "Horizontal dotted lines mark inspection-cycle thresholds at 5 and 30 yr.")},
    {"file": "fig_inspection_threshold.png",
     "caption": ("Figure 9. Fraction of trees with absolute fatigue life below several inspection-cycle "
                 "thresholds (1, 5, 10, 30, 100 yr) as a function of decay ratio. "
                 "Solid lines: lower-bound S-N. Dashed: mean S-N.")},
    {"file": "fig_sens_prior.png",
     "caption": ("Figure 10a. Sensitivity to the Beta prior on normalised eccentricity. "
                 "Four prior shapes (Beta(1,1) uniform, Beta(2,2) baseline, Beta(2,5) low-eccentricity skew, "
                 "Beta(5,2) high-eccentricity skew) at r_d/R = 0.5 and 0.7.")},
    {"file": "fig_sens_aspect.png",
     "caption": ("Figure 10b. Sensitivity to cavity aspect ratio (area-equivalent). "
                 "Circular and elliptical 1:2 / 1:3 cavities at r_d/R = 0.5 and 0.7.")},
    {"file": "fig_sens_variance.png",
     "caption": ("Figure 10c. First-order Sobol-style variance decomposition of the life ratio. "
                 "For circular cavities, eccentricity alone explains 99.6% of the variance; "
                 "for elliptical cavities, eccentricity (71%) plus aspect ratio (8%) "
                 "dominate over orientation variables.")},
    {"file": "fig_policy_matrix.png",
     "caption": ("Figure 11. Risk-based inspection policy matrix (lower-bound S-N, ginkgo). "
                 "Left: recommended inspection interval Delta_t = T_p5 / 2 [yr]. "
                 "Right: categorical risk classification (BI-ANNUAL through REMOVE).")},
    {"file": "fig_abaqus_validation.png",
     "caption": ("Figure A1. Independent ABAQUS validation of the closed-form section properties. "
                 "(a) Per-quantity relative error for six cases (r_d/R = 0.5, 0.7 and "
                 "e_norm = 0.0, 0.5, 0.95). "
                 "(b) Parity plot of all 36 quantities (6 cases x 6 properties). "
                 "Maximum relative error: 7.96e-04.")},
]


# ─────────────────────────────────────────────────────────────────────
# Step 1: Trim MD to body only
# ─────────────────────────────────────────────────────────────────────
print("Step 1: Preparing body-only MD...")
md_text = MD_PATH.read_text(encoding="utf-8")
# Cut at "## List of Figures" -- we re-insert figures at the end with images
cut_match = re.search(r"\n## List of Figures", md_text)
if cut_match:
    md_body = md_text[:cut_match.start()]
else:
    md_body = md_text
TMP_MD.write_text(md_body, encoding="utf-8")


# ─────────────────────────────────────────────────────────────────────
# Step 2: pandoc MD -> DOCX
# ─────────────────────────────────────────────────────────────────────
print("Step 2: Converting body MD to DOCX via pandoc...")
pypandoc.convert_file(
    str(TMP_MD), "docx",
    outputfile=str(TMP_DOCX),
    extra_args=["--wrap=none", "--standalone"],
)


# ─────────────────────────────────────────────────────────────────────
# Step 3: Open body DOCX, append figures, apply formatting
# ─────────────────────────────────────────────────────────────────────
print("Step 3: Building final DOCX...")
doc = Document(str(TMP_DOCX))

# Margins
sec = doc.sections[0]
sec.top_margin    = Cm(2.5)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.0)
sec.right_margin  = Cm(2.5)


def add_para(doc, text="", bold=False, size=12, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, double_space=True, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(after)
    if double_space:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    if text:
        r = p.add_run(text)
        r.font.name = BODY_FONT
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


# ── List of Figures (each figure on its own page-ish area) ──────────────
add_page_break(doc)
add_para(doc, "List of Figures", bold=True, size=14, after=18)

n_added = 0
for fdata in FIGURES:
    img_path = FIG_DIR / fdata["file"]
    if not img_path.exists():
        print(f"  [WARN] Not found: {img_path}")
        continue

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_after = Pt(6)
    p_img.paragraph_format.space_before = Pt(6)
    p_img.add_run().add_picture(str(img_path), width=Cm(15))

    add_para(doc, fdata["caption"], size=11,
             align=WD_ALIGN_PARAGRAPH.LEFT, double_space=False, after=20)
    n_added += 1


# ─────────────────────────────────────────────────────────────────────
# Step 4: Apply font, spacing, color globally
# ─────────────────────────────────────────────────────────────────────
print("Step 4: Applying font & spacing...")


def remove_hr(para):
    """Remove horizontal rule border drawn from Markdown '---' lines."""
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        pBdr = pPr.find(qn("w:pBdr"))
        if pBdr is not None:
            pPr.remove(pBdr)


def fmt_para(para):
    style_name = para.style.name.lower()
    is_heading = "heading" in style_name
    is_title   = "title"   in style_name
    fs = Pt(14) if is_title else Pt(12)
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    remove_hr(para)
    for run in para.runs:
        run.font.name = BODY_FONT
        run.font.size = fs
        run.font.color.rgb = RGBColor(0, 0, 0)
        if is_heading or is_title:
            run.font.bold = True
        rPr = run._r.get_or_add_rPr()
        el = rPr.find(qn("w:rFonts"))
        if el is not None:
            el.set(qn("w:ascii"),  BODY_FONT)
            el.set(qn("w:hAnsi"),  BODY_FONT)
            el.set(qn("w:cs"),     BODY_FONT)
        cel = rPr.find(qn("w:color"))
        if cel is not None:
            cel.attrib.pop(qn("w:themeColor"), None)
            cel.attrib.pop(qn("w:themeShade"), None)
            cel.set(qn("w:val"), "000000")


for para in doc.paragraphs:
    fmt_para(para)

# Tables: single-spaced cells, black text, force black font
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                remove_hr(para)
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.name = BODY_FONT


# ─────────────────────────────────────────────────────────────────────
# Save
# ─────────────────────────────────────────────────────────────────────
doc.save(str(OUT_PATH))
print(f"\nSaved: {OUT_PATH}")

# Cleanup
for f in [TMP_MD, TMP_DOCX]:
    if f.exists():
        f.unlink()

print("Done.")
print(f"  Figures appended : {n_added}")
print(f"  Tables (inline)  : 4 (Tables 1-3 + Table A1, rendered by pandoc)")
